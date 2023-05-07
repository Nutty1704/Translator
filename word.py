from googletrans import Translator
import os
from util import get_language_code
from docx.api import Document as DocumentRead
from docx import Document as DocumentWrite

class WordTranslator:
    def __init__(self) -> None:
        self.translator = Translator()
        self.translate()

    def get_word_file(self) -> str:
        while True:
            file_path = input('Enter the file path: ')
            if os.path.exists(file_path):
                if file_path.endswith('.docx'):
                    return file_path
                else:
                    print('Invalid file extension. Try again.')
            else:
                print('Invalid file path. Try again.')

    def get_language_code(self) -> str:
        while True:
            lang = input('Enter the language: ')
            lang_code = get_language_code(lang)

            if lang_code:
                return lang_code
            else:
                print('Invalid language. Try again.')

    def translate(self) -> None:
        file_path = self.get_word_file()
        lang_code = self.get_language_code()

        # read the file
        doc = DocumentRead(file_path)
        
        new_path = file_path.rstrip('.docx') + f'-{lang_code}.docx'
        new_doc = DocumentWrite()

        # translate the file
        for p in doc.paragraphs:
            translated = self.translator.translate(p.text, dest=lang_code)
            translated_text = translated.text

            new_p = new_doc.add_paragraph(translated_text)
            new_p.style = p.style

        for t in doc.tables:
            new_table = new_doc.add_table(rows=len(t.rows), cols=len(t.columns))
            for i, row in enumerate(t.rows):
                for j, cell in enumerate(row.cells):
                    new_table.cell(i, j).text = self.translator.translate(cell.text, dest=lang_code).text

            new_table.style = t.style
        
        new_doc.save(new_path)
        print('Translation complete.')


if __name__ == "__main__":
    WordTranslator()